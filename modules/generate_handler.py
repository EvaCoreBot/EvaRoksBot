from io import BytesIO
import asyncio
from telegram import Update
from telegram.ext import ContextTypes
from docx import Document
from .gemini import generate_text

async def handle_generate(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    """Generate a DOCX document using Gemini based on the provided topic."""
    topic = " ".join(context.args).strip()
    if not topic:
        await update.message.reply_text("⚠️ Укажите тему: /generate <тема>")
        return

    prompt = f"Составь краткий юридический документ по теме: {topic}."
    try:
        text = await asyncio.to_thread(generate_text, prompt)
    except Exception as exc:
        await update.message.reply_text(f"Ошибка генерации: {exc}")
        return

    document = Document()
    document.add_heading("Сгенерированный документ", 0)
    document.add_paragraph(text)
    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    await update.message.reply_document(document=buffer, filename="generated.docx")
